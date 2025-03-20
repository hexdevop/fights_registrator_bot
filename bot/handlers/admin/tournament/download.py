from aiogram import types, Router, F
from io import BytesIO
from docx import Document
from docx.shared import Pt
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from bot.keyboards.admin.factory import TournamentCallbackData
from database import get_slave_session
from database.models import User
from database.models.admin import Region, Discipline, Tournament

router = Router()


@router.callback_query(TournamentCallbackData.filter(F.action == 'download'))
async def download_docx(
        call: types.CallbackQuery,
        callback_data: TournamentCallbackData,
):
    async with get_slave_session() as session:
        file = await export_tournament_to_docx(session, callback_data.id)
    await call.message.answer_document(
        document=types.BufferedInputFile(file.getvalue(), 'out.docx')
    )


async def export_tournament_to_docx(session: AsyncSession, tournament_id: int) -> BytesIO:
    tournament_stmt = select(Tournament).where(Tournament.id == tournament_id)
    tournament_result = await session.execute(tournament_stmt)
    tournament = tournament_result.scalar_one_or_none()

    if not tournament:
        raise ValueError("Tournament not found")

    stmt = (
        select(User, Discipline.name, Region.name)
        .join(Discipline, User.discipline_id == Discipline.id)
        .join(Region, User.region_id == Region.id)
        .where(User.tournament_id == tournament_id)
        .order_by(User.id)
    )

    result = await session.execute(stmt)
    users = result.all()

    doc = Document()

    doc.add_heading(f"{tournament.name}", level=1)
    doc.add_paragraph(f"Дата проведения: {tournament.date.strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"Организатор: {tournament.organizer}")
    doc.add_paragraph(f"Возрастная категория: {tournament.age}")
    doc.add_paragraph(" ")

    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'ФИО'
    hdr_cells[2].text = 'Дата рождения'
    hdr_cells[3].text = 'Пол'
    hdr_cells[4].text = 'Вес'
    hdr_cells[5].text = 'ФИО Тренера'
    hdr_cells[6].text = 'Регион'
    hdr_cells[7].text = 'Дисциплина'

    for idx, (user, discipline_name, region_name) in enumerate(users, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = user.initials
        row_cells[2].text = user.date
        row_cells[3].text = "Муж" if user.gender == 1 else "Жен"
        row_cells[4].text = user.weight
        row_cells[5].text = user.coach_initials
        row_cells[6].text = region_name
        row_cells[7].text = discipline_name

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream